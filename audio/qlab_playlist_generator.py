#!/usr/bin/env python3
"""
Генерация плейлиста QLab из списка детей и сегментированного аудио (fuzzy matching).
"""

from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

import librosa
import pandas as pd
from thefuzz import fuzz
from tqdm import tqdm

from config import CONFIG, PROJECT_ROOT, SEGMENTED_OUTPUT_DIR

AUDIO_EXTENSIONS = {".wav", ".mp3", ".flac", ".ogg", ".m4a"}


def _applescript_escape_quoted(s: str) -> str:
    """Экранирование для строк в двойных кавычках AppleScript."""
    t = s.replace("\r", " ").replace("\n", " ")
    return t.replace("\\", "\\\\").replace('"', '\\"')


def _write_applescript_file(path: Path, lines: list[str]) -> None:
    """UTF-8 с BOM: кириллица читается верно; без BOM редактор часто берёт MacRoman.
    UTF-16 для .applescript на части систем даёт «неизвестный маркер» при разборе."""
    content = "\n".join(lines)
    if content and not content.endswith("\n"):
        content += "\n"
    path.write_text(content, encoding="utf-8-sig", newline="\n")


@dataclass
class ChildInfo:
    number: int
    surname: str
    name: str
    class_name: str
    notes: str = ""


@dataclass
class PlaylistCueRow:
    """Одна строка плейлиста (один аудиофайл = один кью в QLab)."""

    child_number: int
    surname: str
    name: str
    class_name: str
    notes: str
    matched_file: str = ""
    matched_path: str = ""
    duration: float = 0.0
    match_score: int = 0
    # red — несколько файлов под одного ребёнка (дубли); none — один вариант.
    q_color: str = "none"
    # суффикс к имени кью при нескольких вариантах, напр. " [1/2]"
    cue_name_suffix: str = ""


@dataclass
class AudioSegmentInfo:
    path: Path
    class_folder: str
    label: str
    duration: float


def normalize_header(s: str) -> str:
    return re.sub(r"\s+", " ", str(s).strip().lower())


def _map_header_to_column(h: str) -> str | None:
    """Стандартное имя колонки или None, если не распознано."""
    h = normalize_header(h)
    if h in ("№", "n", "no", "номер", "#", "num", "number", "№ п/п", "п/п"):
        return "number"
    if h.startswith("№") and "п" in h:
        return "number"
    if "фамил" in h or h in ("surname", "last_name", "lastname"):
        return "surname"
    if h in ("имя", "name", "first_name", "firstname", "имя "):
        return "name"
    if "класс" in h or h in ("class", "группа"):
        return "class_name"
    if "примеч" in h or h in ("notes", "note", "коммент", "comment"):
        return "notes"
    return None


# «1А», «10Б» и т.п. — второй символ часто кириллица
_CLASS_TOKEN_RE = re.compile(r"^\d{1,2}[А-ЯЁA-Za-z]?$", re.UNICODE)


def _looks_like_class_token(s: str) -> bool:
    t = s.replace(" ", "").strip()
    return bool(t) and bool(_CLASS_TOKEN_RE.match(t))


def _try_docx_standard_table(table: Any) -> pd.DataFrame | None:
    """Таблица с явными заголовками №, Фамилия, Имя, Класс."""
    if len(table.rows) < 2:
        return None
    header_cells = [c.text for c in table.rows[0].cells]
    col_index: dict[str, int] = {}
    for i, raw in enumerate(header_cells):
        std = _map_header_to_column(raw)
        if std and std not in col_index:
            col_index[std] = i
    if not {"number", "surname", "name", "class_name"}.issubset(col_index.keys()):
        return None
    records: list[dict[str, str]] = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue

        def cell(key: str) -> str:
            j = col_index[key]
            return cells[j] if j < len(cells) else ""

        records.append(
            {
                "number": cell("number"),
                "surname": cell("surname"),
                "name": cell("name"),
                "class_name": cell("class_name"),
                "notes": cell("notes") if "notes" in col_index else "",
            }
        )
    if records:
        return pd.DataFrame(records, dtype=str)
    return None


def _row_looks_like_docx_header(cells: list[str]) -> bool:
    if not cells:
        return False
    h0 = normalize_header(cells[0])
    h1 = normalize_header(cells[1]) if len(cells) > 1 else ""
    if "№" in cells[0] or h0 in ("номер", "n", "#"):
        return True
    if "класс" in h1 or h1 == "class":
        return True
    if "фамил" in h0 or "имя" in normalize_header(cells[2] if len(cells) > 2 else ""):
        return True
    return False


def _try_docx_compact_three_col(table: Any) -> pd.DataFrame | None:
    """
    Три колонки: № (часто только в начале блока), класс, «Имя Фамилия» в одной ячейке.
    """
    if len(table.rows) < 1:
        return None
    if len(table.rows[0].cells) != 3:
        return None

    raw_rows: list[list[str]] = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if not any(cells):
            continue
        raw_rows.append(cells)

    if not raw_rows:
        return None
    if _row_looks_like_docx_header(raw_rows[0]):
        raw_rows = raw_rows[1:]
    if len(raw_rows) < 1:
        return None

    sample = 0
    for cells in raw_rows[:5]:
        if len(cells) < 3:
            continue
        if _looks_like_class_token(cells[1]) and len(cells[2].split()) >= 2:
            sample += 1
    if sample < 1:
        return None

    records: list[dict[str, str]] = []
    seq = 0
    last_klass = ""
    for cells in raw_rows:
        if len(cells) < 3:
            continue
        c0, klass, full = cells[0], cells[1], cells[2]
        if _looks_like_class_token(klass):
            last_klass = klass
        elif not klass.strip() and last_klass:
            klass = last_klass
        if not _looks_like_class_token(klass) or not full:
            continue
        parts = full.split()
        if len(parts) >= 2:
            surname = parts[-1]
            name = " ".join(parts[:-1])
        else:
            name = ""
            surname = parts[0] if parts else ""
        seq += 1
        extra = []
        if c0:
            extra.append(f"№вдок={c0}")
        notes = "; ".join(extra)
        records.append(
            {
                "number": str(seq),
                "surname": surname,
                "name": name,
                "class_name": klass,
                "notes": notes,
            }
        )

    if records:
        return pd.DataFrame(records, dtype=str)
    return None


def _load_children_from_docx(path: Path) -> pd.DataFrame:
    """Таблица в .docx: либо №/Фамилия/Имя/Класс, либо 3 колонки №/Класс/«Имя Фамилия»."""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("Установите пакет: pip install python-docx") from e

    doc = Document(str(path))
    for table in doc.tables:
        df = _try_docx_standard_table(table)
        if df is not None and len(df) > 0:
            return df
    for table in doc.tables:
        df = _try_docx_compact_three_col(table)
        if df is not None and len(df) > 0:
            return df

    raise ValueError(
        "В документе Word не найдена подходящая таблица: нужны колонки "
        "№, Фамилия, Имя, Класс — или три колонки: № (опц.), класс (1А…), "
        "полное имя «Имя Фамилия»."
    )


def load_children_dataframe(file_path: Path) -> pd.DataFrame:
    """Excel, CSV, TXT или Word (.docx) с таблицей: №, Фамилия, Имя, Класс, Примечание."""
    path = Path(file_path)
    if not path.is_file():
        raise FileNotFoundError(f"Файл списка не найден: {path}")

    if path.suffix.lower() == ".docx":
        return _load_children_from_docx(path)
    if path.suffix.lower() in (".xlsx", ".xls"):
        df = pd.read_excel(path, dtype=str)
    elif path.suffix.lower() == ".csv":
        df = pd.read_csv(path, encoding="utf-8-sig", dtype=str)
    elif path.suffix.lower() == ".txt":
        df = pd.read_csv(path, sep=r"\t|,|;", engine="python", encoding="utf-8-sig", dtype=str)
    else:
        raise ValueError(f"Неподдерживаемый формат: {path.suffix}")

    # Нормализация имён колонок
    rename: dict[str, str] = {}
    for c in df.columns:
        std = _map_header_to_column(str(c))
        if std:
            rename[c] = std
    df = df.rename(columns=rename)

    required = {"number", "surname", "name", "class_name"}
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Не хватает колонок: {missing}. Есть: {list(df.columns)}")

    if "notes" not in df.columns:
        df["notes"] = ""

    return df


@dataclass
class QLabPlaylistGenerator:
    children_list_file: Path
    audio_segments_dir: Path
    output_directory: Path
    fuzzy_threshold: int = 80
    export_formats: list[str] = field(
        default_factory=lambda: ["csv", "json", "cue", "applescript"]
    )
    class_filter: str | None = None

    children: list[ChildInfo] = field(default_factory=list)
    audio_segments: list[AudioSegmentInfo] = field(default_factory=list)
    playlist_rows: list[PlaylistCueRow] = field(default_factory=list)

    @classmethod
    def from_config(cls, cfg: dict[str, Any] | None = None) -> QLabPlaylistGenerator:
        c = dict(cfg or CONFIG)
        seg_dir = c.get("audio_segments_dir") or c.get("output_directory") or str(SEGMENTED_OUTPUT_DIR)
        return cls(
            children_list_file=Path(c.get("children_list_file", str(PROJECT_ROOT / "children_list.xlsx"))),
            audio_segments_dir=Path(seg_dir),
            output_directory=Path(c.get("qlab_output_directory", str(PROJECT_ROOT / "qlab_playlist"))),
            fuzzy_threshold=int(c.get("fuzzy_threshold", 80)),
            export_formats=list(c.get("export_formats", ["csv", "json", "cue"])),
            class_filter=c.get("class_filter"),
        )

    def load_children_list(self, file_path: Path | None = None) -> list[ChildInfo]:
        fp = Path(file_path or self.children_list_file)
        df = load_children_dataframe(fp)
        out: list[ChildInfo] = []
        for _, row in df.iterrows():
            try:
                num = int(float(str(row["number"]).replace(",", ".").strip()))
            except (ValueError, TypeError):
                continue
            out.append(
                ChildInfo(
                    number=num,
                    surname=str(row.get("surname", "") or "").strip(),
                    name=str(row.get("name", "") or "").strip(),
                    class_name=str(row.get("class_name", "") or "").strip(),
                    notes=str(row.get("notes", "") or "").strip(),
                )
            )
        self.children = sorted(out, key=lambda x: x.number)
        return self.children

    @staticmethod
    def normalize_name(text: str) -> str:
        t = re.sub(r"\s+", " ", text.strip().lower())
        return t

    @staticmethod
    def filename_to_label(stem: str) -> str:
        """`01_Иван_Петров` → «Иван Петров»."""
        parts = stem.split("_")
        if len(parts) >= 2 and parts[0].isdigit():
            parts = parts[1:]
        return " ".join(p for p in parts if p).strip()

    def scan_audio_segments(self, directory: Path | None = None) -> list[AudioSegmentInfo]:
        root = Path(directory or self.audio_segments_dir)
        if not root.is_dir():
            raise FileNotFoundError(f"Каталог сегментов не найден: {root}")

        segments: list[AudioSegmentInfo] = []
        for fpath in sorted(root.rglob("*")):
            if not fpath.is_file():
                continue
            if fpath.name.startswith("processing_report"):
                continue
            if fpath.suffix.lower() not in AUDIO_EXTENSIONS:
                continue
            class_folder = fpath.parent.name
            label = self.filename_to_label(fpath.stem)
            dur = self.get_audio_duration(fpath)
            segments.append(
                AudioSegmentInfo(
                    path=fpath.resolve(),
                    class_folder=class_folder,
                    label=label,
                    duration=dur,
                )
            )

        self.audio_segments = segments
        return segments

    @staticmethod
    def get_audio_duration(file_path: Path | str) -> float:
        p = Path(file_path)
        try:
            d = librosa.get_duration(path=str(p))
            return float(d)
        except Exception:
            return 0.0

    @staticmethod
    def calculate_match_score(name1: str, name2: str) -> int:
        a = QLabPlaylistGenerator.normalize_name(name1)
        b = QLabPlaylistGenerator.normalize_name(name2)
        return int(
            max(
                fuzz.token_sort_ratio(a, b),
                fuzz.partial_ratio(a, b),
            )
        )

    def _class_matches(self, child_class: str, folder: str) -> bool:
        if self.class_filter is not None and str(self.class_filter).strip() != "":
            if self.normalize_name(folder) != self.normalize_name(str(self.class_filter)):
                return False
        c1 = self.normalize_name(child_class)
        f1 = self.normalize_name(folder)
        if c1 == f1:
            return True
        if c1 and f1 and (c1 in f1 or f1 in c1):
            return True
        return fuzz.ratio(c1, f1) >= 70

    def match_names(
        self,
        threshold: int | None = None,
    ) -> list[PlaylistCueRow]:
        th = threshold if threshold is not None else self.fuzzy_threshold
        rows: list[PlaylistCueRow] = []

        for ch in self.children:
            child_variants = [
                f"{ch.name} {ch.surname}",
                f"{ch.surname} {ch.name}",
            ]
            child_variants = [v.strip() for v in child_variants if v.strip()]

            candidates: list[tuple[AudioSegmentInfo, int]] = []
            for seg in self.audio_segments:
                if not self._class_matches(ch.class_name, seg.class_folder):
                    continue
                best_for_seg = 0
                for cv in child_variants:
                    sc = self.calculate_match_score(cv, seg.label)
                    if sc > best_for_seg:
                        best_for_seg = sc
                if best_for_seg >= th:
                    candidates.append((seg, best_for_seg))

            candidates.sort(key=lambda x: (-x[1], str(x[0].path)))
            seen_paths: set[str] = set()
            uniq: list[tuple[AudioSegmentInfo, int]] = []
            for seg, sc in candidates:
                key = str(seg.path.resolve())
                if key in seen_paths:
                    continue
                seen_paths.add(key)
                uniq.append((seg, sc))

            dup = len(uniq) > 1
            qcol = "red" if dup else "none"

            if not uniq:
                best_score = 0
                for seg in self.audio_segments:
                    if not self._class_matches(ch.class_name, seg.class_folder):
                        continue
                    for cv in child_variants:
                        best_score = max(
                            best_score, self.calculate_match_score(cv, seg.label)
                        )
                rows.append(
                    PlaylistCueRow(
                        child_number=ch.number,
                        surname=ch.surname,
                        name=ch.name,
                        class_name=ch.class_name,
                        notes=ch.notes,
                        match_score=best_score,
                        q_color="none",
                    )
                )
            else:
                for ki, (seg, sc) in enumerate(uniq, start=1):
                    suffix = f" [{ki}/{len(uniq)}]" if len(uniq) > 1 else ""
                    rows.append(
                        PlaylistCueRow(
                            child_number=ch.number,
                            surname=ch.surname,
                            name=ch.name,
                            class_name=ch.class_name,
                            notes=ch.notes,
                            matched_file=seg.path.name,
                            matched_path=str(seg.path.resolve()),
                            duration=seg.duration,
                            match_score=sc,
                            q_color=qcol,
                            cue_name_suffix=suffix,
                        )
                    )

        self.playlist_rows = rows
        return rows

    def generate_playlist(self) -> list[PlaylistCueRow]:
        self.load_children_list()
        self.scan_audio_segments()
        return self.match_names()

    @staticmethod
    def _cue_display_name(row: PlaylistCueRow) -> str:
        return (f"{row.name} {row.surname}".strip() + row.cue_name_suffix).strip()

    def export_csv(self, output_path: Path | str) -> Path:
        p = Path(output_path)
        p.parent.mkdir(parents=True, exist_ok=True)
        with open(p, "w", newline="", encoding="utf-8") as f:
            w = csv.writer(f)
            w.writerow(
                ["Cue Number", "Cue Name", "File Path", "Duration", "Notes", "Q Color"]
            )
            for i, row in enumerate(self.playlist_rows, start=1):
                name = self._cue_display_name(row)
                if not row.matched_path:
                    fp, dur = "", ""
                else:
                    fp, dur = row.matched_path, f"{row.duration:.2f}"
                notes_parts = [row.class_name]
                if row.notes:
                    notes_parts.append(row.notes)
                notes = " ".join(notes_parts)
                qcol = row.q_color if row.q_color != "none" else ""
                w.writerow([i, name, fp, dur, notes, qcol])
        return p

    def export_json(self, output_path: Path | str) -> Path:
        p = Path(output_path)
        p.parent.mkdir(parents=True, exist_ok=True)
        cues = []
        for i, row in enumerate(self.playlist_rows, start=1):
            name = self._cue_display_name(row)
            notes = row.class_name + (f" {row.notes}" if row.notes else "")
            cues.append(
                {
                    "number": str(i),
                    "name": name,
                    "type": "Audio",
                    "file": row.matched_path or "",
                    "duration": round(row.duration, 2) if row.matched_path else 0.0,
                    "notes": notes.strip(),
                    "qColor": row.q_color,
                }
            )
        doc = {"workspace": {"version": "5.0", "cues": cues}}
        p.write_text(json.dumps(doc, ensure_ascii=False, indent=2), encoding="utf-8")
        return p

    def export_cue(self, output_path: Path | str) -> Path:
        p = Path(output_path)
        p.parent.mkdir(parents=True, exist_ok=True)
        lines = [
            'REM GENRE "School Performance"',
            f'REM DATE "{datetime.now().year}"',
            "",
        ]
        for i, row in enumerate(self.playlist_rows, start=1):
            name = self._cue_display_name(row)
            if row.matched_path:
                ap = str(Path(row.matched_path).resolve())
                lines.append(f'FILE "{ap}" WAVE')
            else:
                lines.append('FILE "" WAVE')
            lines.append(f"  TRACK {i:02d} AUDIO")
            lines.append(f'    TITLE "{name}"')
            lines.append("    INDEX 01 00:00:00")
            lines.append("")
        p.write_text("\n".join(lines), encoding="utf-8")
        return p

    def export_applescript(self, output_path: Path | str) -> Path:
        """Скрипт для QLab 5: Audio cues (make + selected).

        Кириллица внутри кавычек в теле .applescript ломает компилятор Редактора сценариев.
        Пути и имена выносятся в UTF-8 sidecar-файлы (.paths.txt, .names.txt, .colors.txt);
        сценарий читает их через «cat» (в исходнике только ASCII).
        Имя кью: только «set q name» — «q list name» в словаре QLab только для чтения, а слово list
        в AppleScript зарезервировано и даёт синтаксическую ошибку при компиляции.
        """
        p = Path(output_path)
        p.parent.mkdir(parents=True, exist_ok=True)
        qlab_id = "com.figure53.qlab.5"
        stem = p.stem
        paths_sidecar = p.parent / f"{stem}.paths.txt"
        names_sidecar = p.parent / f"{stem}.names.txt"
        colors_sidecar = p.parent / f"{stem}.colors.txt"

        rows_with_file = [r for r in self.playlist_rows if r.matched_path]
        if not rows_with_file:
            lines = [
                "-- QLab 5: open workspace, click inside the target cue list.",
                "-- No cues: sidecar files not used.",
                f"-- {datetime.now().isoformat(timespec='seconds')}",
                "",
                f'tell application id "{qlab_id}"',
                "  activate",
                "  delay 0.4",
                "  tell front workspace",
                '    set emptyMsg to "No audio file paths in playlist. Match children to segments first."',
                '    display dialog emptyMsg buttons {"OK"} default button 1 with title "QLab"',
                "  end tell",
                "end tell",
                "",
            ]
            _write_applescript_file(p, lines)
            for side in (paths_sidecar, names_sidecar, colors_sidecar):
                if side.exists():
                    side.unlink()
            return p

        path_lines: list[str] = []
        name_lines: list[str] = []
        color_lines: list[str] = []
        for row in rows_with_file:
            path_lines.append(str(Path(row.matched_path).resolve()))
            nm = self._cue_display_name(row).replace("\r", " ").replace("\n", " ")
            name_lines.append(nm)
            color_lines.append(row.q_color if row.q_color in ("red", "none") else "none")

        paths_sidecar.write_text("\n".join(path_lines) + "\n", encoding="utf-8")
        names_sidecar.write_text("\n".join(name_lines) + "\n", encoding="utf-8")
        colors_sidecar.write_text("\n".join(color_lines) + "\n", encoding="utf-8")

        pf = _applescript_escape_quoted(str(paths_sidecar.resolve()))
        nf = _applescript_escape_quoted(str(names_sidecar.resolve()))
        cf = _applescript_escape_quoted(str(colors_sidecar.resolve()))

        lines = [
            "-- QLab 5: keep this .applescript next to .paths.txt, .names.txt, .colors.txt (UTF-8).",
            "-- Open workspace, click in cue list, then Run in Script Editor.",
            f"-- {datetime.now().isoformat(timespec='seconds')}",
            "",
            f'tell application id "{qlab_id}"',
            "  activate",
            "  delay 0.4",
            "  tell front workspace",
            f'    set pathsFile to "{pf}"',
            f'    set namesFile to "{nf}"',
            f'    set colorsFile to "{cf}"',
            '    set rawPaths to do shell script "cat " & quoted form of pathsFile',
            '    set rawNames to do shell script "cat " & quoted form of namesFile',
            '    set rawColors to do shell script "cat " & quoted form of colorsFile',
            "    set pathLines to paragraphs of rawPaths",
            "    set nameLines to paragraphs of rawNames",
            "    set colorLines to paragraphs of rawColors",
            "    set np to count of pathLines",
            "    set nn to count of nameLines",
            "    set nc to count of colorLines",
            "    if np is not equal to nn or np is not equal to nc then",
            '      display dialog "Sidecar line count mismatch (paths/names/colors)." buttons {"OK"} default button 1 with title "QLab"',
            "    else",
            "      set failCount to 0",
            "      repeat with i from 1 to np",
            "        set onePath to item i of pathLines",
            "        if (length of onePath) > 0 then",
            "          try",
            '            make type "audio"',
            "            set newCue to last item of (selected as list)",
            "            set file target of newCue to POSIX file onePath",
            "            set q name of newCue to item i of nameLines",
            "            set q color of newCue to item i of colorLines",
            "          on error errMsg",
            '            log "QLab cue " & i & ": " & errMsg',
            "            set failCount to failCount + 1",
            "          end try",
            "          delay 0.05",
            "        end if",
            "      end repeat",
            '      if failCount > 0 then log "QLab: failed " & failCount & " cue(s) (see log above)"',
            "    end if",
            "  end tell",
            "end tell",
            "",
        ]
        _write_applescript_file(p, lines)
        return p

    def generate_report(self, output_path: Path | str) -> Path:
        p = Path(output_path)
        p.parent.mkdir(parents=True, exist_ok=True)
        by_child: dict[int, list[PlaylistCueRow]] = {}
        for row in self.playlist_rows:
            by_child.setdefault(row.child_number, []).append(row)

        n_cues = sum(1 for r in self.playlist_rows if r.matched_path)
        n_children_ok = sum(
            1
            for cn in by_child
            if any(r.matched_path for r in by_child[cn])
        )
        n_children_missing = len(self.children) - n_children_ok

        lines = [
            "=" * 60,
            "ОТЧЁТ СОПОСТАВЛЕНИЯ АУДИО И СПИСКА ДЕТЕЙ",
            f"Создан: {datetime.now().isoformat(timespec='seconds')}",
            "=" * 60,
            f"Детей в списке: {len(self.children)}",
            f"Строк в плейлисте (кью): {len(self.playlist_rows)}",
            f"С файлами: {n_cues} кью, детей с хотя бы одним файлом: {n_children_ok}",
            f"Без ни одного файла: {n_children_missing}",
            "",
            "Детали:",
            "-" * 60,
        ]
        for cn in sorted(by_child.keys()):
            rows = by_child[cn]
            name = f"{rows[0].surname} {rows[0].name}"
            if any(r.matched_path for r in rows):
                for r in rows:
                    if not r.matched_path:
                        continue
                    dup = " [дубли — красный в QLab]" if r.q_color == "red" else ""
                    lines.append(
                        f"  ✓ №{cn}: {name} → {r.matched_file} ({r.match_score}%){dup}"
                    )
            else:
                lines.append(
                    f"  ⚠ №{cn}: {name} → НЕ НАЙДЕН (лучший балл: {rows[0].match_score})"
                )
        p.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return p


def main() -> int:
    cfg = dict(CONFIG)
    tqdm.write("")
    tqdm.write("=" * 60)
    tqdm.write("🎵 ГЕНЕРАТОР ПЛЕЙЛИСТА ДЛЯ QLAB")
    tqdm.write("=" * 60)

    gen = QLabPlaylistGenerator.from_config(cfg)

    tqdm.write("📖 Чтение списка детей...")
    try:
        gen.load_children_list()
    except Exception as e:
        tqdm.write(f"❌ Ошибка чтения списка: {e}")
        return 1
    tqdm.write(f"✅ Найдено детей: {len(gen.children)}")

    tqdm.write("🔍 Сканирование аудиофайлов...")
    try:
        gen.scan_audio_segments()
    except Exception as e:
        tqdm.write(f"❌ Ошибка сканирования: {e}")
        return 1
    tqdm.write(f"✅ Найдено сегментов: {len(gen.audio_segments)}")

    tqdm.write(f"🔗 Сопоставление имён (порог: {gen.fuzzy_threshold}%)...")
    tqdm.write("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    gen.match_names()
    for row in gen.playlist_rows:
        name = gen._cue_display_name(row)
        if row.matched_path:
            dup = " [дубль→красный]" if row.q_color == "red" else ""
            tqdm.write(
                f"  ✓ №{row.child_number}: {name} → {row.matched_file} ({row.match_score}%){dup}"
            )
        else:
            tqdm.write(f"  ⚠ №{row.child_number}: {name} → НЕ НАЙДЕН")

    n_ok = sum(1 for c in gen.children if any(
        r.child_number == c.number and r.matched_path for r in gen.playlist_rows
    ))
    n_missing = len(gen.children) - n_ok
    tqdm.write("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━")
    tqdm.write("📊 Статистика сопоставления:")
    tqdm.write(f"   ✅ Детей с файлом: {n_ok}")
    tqdm.write(f"   ⚠ Без файла: {n_missing}")

    out_dir = gen.output_directory
    out_dir.mkdir(parents=True, exist_ok=True)

    tqdm.write("💾 Экспорт плейлистов...")
    fmt = [f.lower() for f in gen.export_formats]
    if "csv" in fmt:
        path = gen.export_csv(out_dir / "playlist.csv")
        tqdm.write(f"   ✓ CSV: {path}")
    if "json" in fmt:
        path = gen.export_json(out_dir / "qlab_workspace.json")
        tqdm.write(f"   ✓ JSON: {path}")
    if "cue" in fmt:
        path = gen.export_cue(out_dir / "playlist.cue")
        tqdm.write(f"   ✓ CUE: {path}")
    if "applescript" in fmt:
        path = gen.export_applescript(out_dir / "build_qlab_playlist.applescript")
        tqdm.write(f"   ✓ AppleScript (QLab 5): {path}")

    gen.generate_report(out_dir / "matching_report.txt")
    tqdm.write(f"   ✓ Отчёт: {out_dir / 'matching_report.txt'}")

    tqdm.write("=" * 60)
    tqdm.write("✅ ГОТОВО! CSV / CUE или build_qlab_playlist.applescript в QLab 5")
    tqdm.write("=" * 60)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
